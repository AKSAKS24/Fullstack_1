"""
Utilities for building DOCX documents from generated sections.  Uses the
`python-docx` library to assemble a structured document with headings and
paragraphs.  Table support is rudimentary; formatting options can be extended
via the `formatting` parameter.
"""
from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import Inches


class DocxBuilder:
    def __init__(self, formatting: Optional[Dict[str, Any]] = None) -> None:
        self.formatting = formatting or {}

    def build(self, sections: List[Dict[str, Any]], output_path: str) -> None:
        """
        Build a DOCX file from a list of sections.  Each section dict should
        contain `title` and `content`.  If the section content is a list of
        dictionaries, a simple table is generated instead of plain text.
        """
        doc = Document()
        for section in sections:
            title = section.get("title", "Section")
            content = section.get("content", "")
            doc.add_heading(title, level=2)
            if isinstance(content, list):
                # build table
                if content:
                    headers = list(content[0].keys())
                    table = doc.add_table(rows=1, cols=len(headers))
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(headers):
                        hdr_cells[i].text = str(h)
                    for row in content:
                        row_cells = table.add_row().cells
                        for i, h in enumerate(headers):
                            row_cells[i].text = str(row.get(h, ""))
            else:
                for para in str(content).split("\n"):
                    doc.add_paragraph(para)
        doc.save(output_path)